# Файл: services/file_processor.py

import logging
from typing import Iterable

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

TXT_ENCODINGS = ("utf-8-sig", "utf-8", "cp1251")
PDF_TEXT_SOFT_LIMIT = 60_001


def _join_clean_text_blocks(blocks: Iterable[str], separator: str = "\n") -> str:
    """
    Об'єднує текстові блоки, прибираючи порожні значення.
    Це допомагає не створювати зайві пусті рядки після обробки документів.
    """
    clean_blocks = []

    for block in blocks:
        if not block:
            continue

        clean_block = block.strip()
        if clean_block:
            clean_blocks.append(clean_block)

    return separator.join(clean_blocks).strip()


def process_docx(file_path: str) -> str:
    """
    Витягує текст із DOCX-файлу.

    Обробляє:
    - звичайні абзаци;
    - таблиці.
    """
    try:
        doc = Document(file_path)

        text_blocks = []

        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_blocks.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                cells_text = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text and cell.text.strip()
                ]

                if cells_text:
                    text_blocks.append(" | ".join(cells_text))

        text = _join_clean_text_blocks(text_blocks)

        if not text:
            logger.warning("DOCX файл не містить тексту або текст не вдалося витягнути: %s", file_path)

        return text

    except Exception:
        logger.exception("Помилка читання DOCX-файлу: %s", file_path)
        raise


def process_pdf(file_path: str) -> str:
    """
    Витягує текст із PDF-файлу.

    PDF-документ явно закривається після обробки,
    щоб не тримати файлові ресурси відкритими.
    """
    doc = None

    try:
        doc = fitz.open(file_path)

        text_blocks = []
        extracted_length = 0

        for page in doc:
            page_text = page.get_text()
            if page_text and page_text.strip():
                remaining_length = PDF_TEXT_SOFT_LIMIT - extracted_length

                if remaining_length <= 0:
                    break

                limited_page_text = page_text[:remaining_length]
                text_blocks.append(limited_page_text)
                extracted_length += len(limited_page_text)

                if extracted_length >= PDF_TEXT_SOFT_LIMIT:
                    logger.info(
                        "PDF text extraction soft limit reached for file: %s",
                        file_path,
                    )
                    break

        text = _join_clean_text_blocks(text_blocks, separator="\n\n")

        if not text:
            logger.warning("PDF файл не містить тексту або текст не вдалося витягнути: %s", file_path)

        return text

    except Exception:
        logger.exception("Помилка читання PDF-файлу: %s", file_path)
        raise

    finally:
        if doc is not None:
            doc.close()


def process_txt(file_path: str) -> str:
    """
    Витягує текст із TXT-файлу.

    Спочатку пробує UTF-8 з BOM, потім звичайний UTF-8,
    потім CP1251 для старих кириличних файлів.
    """
    last_error = None

    for encoding in TXT_ENCODINGS:
        try:
            with open(file_path, "r", encoding=encoding) as file:
                text = file.read().strip()

            if not text:
                logger.warning("TXT файл порожній або не містить тексту: %s", file_path)

            return text

        except UnicodeDecodeError as error:
            last_error = error
            logger.warning(
                "Не вдалося прочитати TXT як %s: %s",
                encoding,
                file_path
            )

    logger.exception(
        "Не вдалося прочитати TXT-файл жодним із підтримуваних кодувань: %s",
        file_path
    )

    if last_error:
        raise last_error

    raise UnicodeDecodeError(
        "unknown",
        b"",
        0,
        1,
        "Не вдалося визначити кодування TXT-файлу"
    )
